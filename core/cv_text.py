"""
cv_text.py — Turn a CV file into clean, parseable plain text.

Extraction quality is load-bearing: every downstream stage (skills, titles,
years, generated documents) reads this text, so garbage here is garbage that
reaches an employer's inbox.

Two real-world problems this handles:

  Overlapping text layers. Some PDFs draw two versions of a line at nearly
  the same height — an edit layered over the original. pdfplumber's default
  y_tolerance of 3 merges those baselines into one line and sorts the glyphs
  by x, interleaving two strings into mush. A tight y_tolerance keeps them as
  separate lines. The cost is that small-caps headings split across lines, but
  those are page furniture we strip anyway.

  Private-use bullets. Word exports Wingdings bullets into the Unicode
  private-use area (U+F0A7, U+F0FC, …). They carry no meaning to anything
  downstream, so they're normalised to a plain "• ".
"""

from __future__ import annotations

import logging
import re
from collections import Counter
from pathlib import Path

logger = logging.getLogger(__name__)

# Wingdings/Symbol glyphs Word emits into the private-use area. Anything in
# U+F000–U+F0FF is decorative in a CV, so map the common bullets to "•" and
# drop the rest.
_PUA_BULLETS = {"", "", "", "", "", "", ""}
_PUA_RANGE = re.compile(r"[-]")

_RULE_RE = re.compile(r"^[\s_\-=•·—–]{6,}$")
_PAGE_NUM_RE = re.compile(r"^\s*(page\s+\d+|\d+\s*/\s*\d+|\d{1,2})\s*$", re.I)


def _normalise_glyphs(text: str) -> str:
    for bullet in _PUA_BULLETS:
        text = text.replace(bullet, "• ")
    text = _PUA_RANGE.sub(" ", text)
    return text.replace(" ", " ").replace("’", "'")


# Words that mark a line as a section heading. A CV running over several pages
# often repeats its section heading as a continuation marker, which makes the
# heading look exactly like page furniture — and dropping it costs the whole
# section, since the parser locates sections by their headings. Structure wins
# over tidiness: a heading is never furniture.
_HEADING_WORDS = {
    "experience", "experiences", "employment", "history", "education",
    "skills", "competencies", "projects", "certifications", "certificates",
    "summary", "profile", "objective", "achievements", "publications",
    "references", "awards", "training", "qualifications", "portfolio",
}


def _is_section_heading(line: str) -> bool:
    """
    Whether a line reads as a section heading rather than body text.

    Deliberately loose: a false positive keeps one extra line of page
    furniture, while a false negative deletes an entire CV section.
    """
    bare = line.strip().strip(":")
    if not bare or len(bare) > 45 or re.search(r"\d", bare):
        return False
    words = re.sub(r"[^a-z& ]", " ", bare.lower()).split()
    return bool(words) and any(w in _HEADING_WORDS for w in words)


def _furniture_key(line: str) -> str:
    """
    Normalise a line so the same running header matches across pages.

    Page furniture nearly always carries the page number — "Jane Doe Page 2",
    "Page 3 of 7" — so comparing raw text never finds it: every page's header
    is unique. Blanking the digits is what makes them comparable.
    """
    return re.sub(r"\d+", "#", line).strip().lower()


def _strip_repeated_lines(pages: list[str]) -> list[str]:
    """
    Drop running headers and footers. A short line appearing at the top or
    bottom of most pages is page furniture, not content.

    This matters well beyond tidiness. A header left in the body gets glued
    onto whatever bullet precedes it, and that bullet is quoted verbatim into
    a generated CV — so an employer reads "…back-end systems. Jane Doe Page 3
    Utilized version control…".
    """
    if len(pages) < 2:
        return pages

    counts: Counter[str] = Counter()
    for page in pages:
        lines = [ln.strip() for ln in page.split("\n") if ln.strip()]
        # Count each line at most once per page. On a short page the head and
        # tail slices overlap, and double-counting one line from one page
        # makes it look like it repeats across pages — which would delete
        # real content on any CV with a sparse page.
        candidates = {
            _furniture_key(line)
            for line in lines[:2] + lines[-2:]
            if len(line) <= 60 and not _is_section_heading(line)
        }
        counts.update(candidates)

    threshold = 2 if len(pages) == 2 else max(2, len(pages) // 2)
    furniture = {key for key, n in counts.items() if n >= threshold and key}
    if furniture:
        logger.info(f"[CV] Dropping page furniture: {sorted(furniture)[:4]}")

    cleaned: list[str] = []
    for page in pages:
        kept = [
            ln for ln in page.split("\n")
            if _is_section_heading(ln.strip())
            or _furniture_key(ln.strip()) not in furniture
        ]
        cleaned.append("\n".join(kept))
    return cleaned


def _clean(text: str) -> str:
    text = _normalise_glyphs(text)
    out: list[str] = []
    for line in text.split("\n"):
        stripped = line.rstrip()
        if _RULE_RE.match(stripped.strip()):
            continue
        if _PAGE_NUM_RE.match(stripped.strip()):
            continue
        stripped = re.sub(r"[ \t]{2,}", "  ", stripped)
        out.append(stripped)

    # Collapse runs of blank lines to at most one.
    result = "\n".join(out)
    return re.sub(r"\n{3,}", "\n\n", result).strip()


def _extract_pdf(path: str) -> str:
    import pdfplumber  # type: ignore

    pages: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            # y_tolerance=1 keeps overlapping text layers on separate lines
            # instead of interleaving them into unusable text.
            #
            # x_tolerance stays at pdfplumber's default of 3: the overlap
            # problem is purely vertical, and tightening x splits
            # letter-spaced headings — a name set with wide tracking comes out
            # as "A M O S . O . A D E W O P O" and stops being a name.
            text = page.extract_text(y_tolerance=1) or ""
            pages.append(text)

    pages = _strip_repeated_lines(pages)
    return "\n".join(pages)


def _extract_docx(path: str) -> str:
    from docx import Document  # type: ignore

    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    # Skills and experience are often laid out in tables, which
    # doc.paragraphs skips entirely.
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append("  ".join(cells))
    return "\n".join(parts)


def extract_cv_text(cv_path: str) -> str:
    """Extract clean plain text from a .pdf, .docx, or .txt CV."""
    ext = Path(cv_path).suffix.lower()
    try:
        if ext == ".docx":
            raw = _extract_docx(cv_path)
        elif ext == ".pdf":
            raw = _extract_pdf(cv_path)
        else:
            raw = Path(cv_path).read_text(encoding="utf-8", errors="ignore")
    except Exception as e:
        logger.error(f"[CV] Extraction failed for {cv_path}: {e}")
        return ""

    cleaned = _clean(raw)
    if not cleaned:
        logger.error(f"[CV] No text extracted from {cv_path}")
    else:
        logger.info(f"[CV] Extracted {len(cleaned)} chars from {Path(cv_path).name}")
    return cleaned


def find_emails(text: str) -> list[str]:
    """All distinct email addresses in the text, in order of appearance."""
    seen: list[str] = []
    for match in re.finditer(r"\b[\w.+-]+@[\w-]+\.[\w.-]+\b", text):
        email = match.group(0).lower().rstrip(".,;:")
        if email not in seen:
            seen.append(email)
    return seen


_PHONE_CANDIDATE_RE = re.compile(
    r"(?:\+\d{1,3}[\s.\-]?)?(?:\(\d{1,4}\)[\s.\-]?)?\d[\d\s.\-]{6,14}\d"
)


def _phone_digits(raw: str) -> str:
    return re.sub(r"\D", "", raw)


def find_phones(text: str) -> list[str]:
    """
    All distinct phone numbers in the text.

    Filtered by digit count, which is what separates a phone number from the
    other numbers a CV is full of: a date range like "2018 - 2023" is eight
    digits, a real number is at least nine.
    """
    seen: list[str] = []
    seen_digits: set[str] = set()
    for match in _PHONE_CANDIDATE_RE.finditer(text or ""):
        raw = match.group(0).strip(" .-")
        digits = _phone_digits(raw)
        if not (9 <= len(digits) <= 15):
            continue
        # "(409)655-2769" and "+1 409 655 2769" are the same number written
        # two ways; compare on the last 10 digits to collapse them.
        key = digits[-10:]
        if key in seen_digits:
            continue
        seen_digits.add(key)
        seen.append(raw)
    return seen


def find_all(pattern: re.Pattern, text: str) -> list[str]:
    """All distinct matches of a pattern, in order of appearance."""
    seen: list[str] = []
    for match in pattern.finditer(text or ""):
        value = match.group(0)
        if value not in seen:
            seen.append(value)
    return seen
