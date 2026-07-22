import os
import re
import sys
import logging
import subprocess
from pathlib import Path
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from core.groq_client import chat_json
from core.cv_profile import profile_is_sendable
from core.cv_text import extract_cv_text  # re-exported for existing callers
from core.tailor import (
    build_cv_data,
    build_cover_letter,
    build_email,
    validate_output,
)
from config import config

logger = logging.getLogger(__name__)

# ──────────────────────────────────────────────────────────
# Theme
# ──────────────────────────────────────────────────────────
ACCENT  = RGBColor(0x1A, 0x56, 0xDB)
DARK    = RGBColor(0x0F, 0x17, 0x2A)
MID     = RGBColor(0x37, 0x41, 0x51)
GRAY    = RGBColor(0x6B, 0x72, 0x80)
RULE_BLUE  = "1A56DB"
RULE_LIGHT = "DBEAFE"


# ──────────────────────────────────────────────────────────
# Polish prompts
#
# These rewrite prose that already exists and is already accurate. They are
# not asked to supply facts, because a model asked for facts invents them —
# which is how the previous version ended up with placeholder years and
# renamed projects. Rewriting only what's given keeps the output tied to the
# CV even when the model is having a bad day.
# ──────────────────────────────────────────────────────────

_POLISH_CL_SYSTEM = """You rewrite cover letter paragraphs to read more naturally.

You will be given four paragraphs assembled from a real CV. Rewrite them so
they flow like a person wrote them, not a template.

HARD RULES:
- Use ONLY facts present in the input. Never add a company, technology,
  metric, date, or achievement that is not already there.
- Never invent numbers. If the input has no metric, the output has none.
- Do not start any paragraph with "Dear" or any salutation.
- Keep it under 350 words total.
- Plain language. No "passionate", "thrilled", "excited", "leverage",
  "robust", "seamless", "synergy".

Return ONLY valid JSON:
{"opening_paragraph":"...","body_paragraph_1":"...","body_paragraph_2":"...","closing_paragraph":"..."}"""

_POLISH_EMAIL_SYSTEM = """You rewrite a job application email to sound human.

You will be given a subject and body assembled from a real CV. Rewrite them
to read like a working engineer wrote it in five minutes.

HARD RULES:
- Use ONLY facts present in the input. Invent nothing — no companies, no
  technologies, no metrics, no dates.
- Keep the sign-off block (name, email, phone, link) exactly as given.
- Body 110-170 words. No bullet points. No emojis.
- Avoid "passionate", "thrilled", "excited", "leverage", "robust".

Return ONLY valid JSON:
{"subject":"...","body":"..."}"""


# ──────────────────────────────────────────────────────────
# Main Entry
# ──────────────────────────────────────────────────────────

def _polish_threshold() -> int:
    """
    Score at or above which a job earns an LLM polish pass. Set
    LLM_POLISH_MIN_SCORE=101 to disable polishing entirely and run fully
    offline.
    """
    try:
        return int(getattr(config, "LLM_POLISH_MIN_SCORE", 85))
    except (TypeError, ValueError):
        return 85


def generate_documents(job: dict, profile: dict, contact: dict,
                       score_data: dict, output_dir: str) -> tuple[bool, str]:
    """
    Build the application package for one job.

    Content is assembled offline from the parsed CV profile. Only jobs
    scoring at or above LLM_POLISH_MIN_SCORE get an LLM pass, and that pass
    rewrites prose the deterministic path already produced — so a failed or
    rate-limited call costs nothing but polish.

    Returns (success, message).
    """
    company = _safe_dirname(job.get("company", "Company"))
    role    = _safe_dirname(job.get("title", "Role"))
    folder  = Path(output_dir) / f"{company}_{role}"

    ok, blockers = profile_is_sendable(profile)
    if not ok:
        msg = f"CV profile is not usable: {'; '.join(blockers)}"
        logger.error(f"[DocGen] Refusing to generate for {company} — {msg}")
        return False, msg

    # ── Offline assembly — no API calls ──────────────────────
    cv_data    = build_cv_data(profile, job, score_data)
    cl_data    = build_cover_letter(profile, job, contact, score_data)
    email_data = build_email(profile, job, contact, score_data)

    # ── Optional polish for jobs worth the tokens ────────────
    score = int(score_data.get("score", 0) or 0)
    threshold = _polish_threshold()
    if score >= threshold:
        logger.info(f"[DocGen] {company}: score {score} ≥ {threshold} — polishing")
        cl_data    = _polish_cover_letter(cl_data, profile, job, contact) or cl_data
        email_data = _polish_email(email_data, profile, job, contact) or email_data
    else:
        logger.debug(f"[DocGen] {company}: score {score} < {threshold} — offline only")

    # ── Last line of defence before anything reaches an inbox ─
    problems = validate_output(cv_data, cl_data, email_data)
    if problems:
        msg = f"generated content failed validation: {'; '.join(problems[:3])}"
        logger.error(f"[DocGen] Refusing to write {company} — {msg}")
        return False, msg

    folder.mkdir(parents=True, exist_ok=True)
    cv_docx = folder / "CV.docx"
    cl_docx = folder / "CoverLetter.docx"

    _write_cv(cv_data, str(cv_docx))
    _write_cl(cl_data, job, contact, str(cl_docx), cv_data.get("identity"))
    _write_email(email_data, job, contact, score_data, folder)

    _to_pdf(str(cv_docx), str(folder))
    _to_pdf(str(cl_docx), str(folder))

    logger.info(f"[DocGen] Done → {folder}")
    return True, str(folder)


# ──────────────────────────────────────────────────────────
# Polish passes — high-scoring jobs only
# ──────────────────────────────────────────────────────────

def _fact_guard(original: str, rewritten: str) -> bool:
    """
    Reject a rewrite that introduces numbers the original didn't have.

    A polish pass is only allowed to change wording. New digits mean the
    model invented a metric or a date, which is the exact failure that put
    "YYYY" and fabricated achievements into the old generated CVs.
    """
    original_nums = set(re.findall(r"\d+", original))
    new_nums = set(re.findall(r"\d+", rewritten))
    invented = new_nums - original_nums
    if invented:
        logger.warning(
            f"[Polish] Rejected rewrite — invented numbers: {sorted(invented)[:5]}"
        )
        return False
    return True


def _polish_cover_letter(cl_data: dict, profile: dict, job: dict,
                         contact: dict) -> Optional[dict]:
    keys = ["opening_paragraph", "body_paragraph_1",
            "body_paragraph_2", "closing_paragraph"]
    original = "\n\n".join(cl_data.get(k, "") for k in keys)

    user = (
        f"ROLE: {job.get('title', '')} at {job.get('company', '')}\n"
        f"ADDRESS TO: {contact.get('hr_name') or 'Hiring Manager'}\n\n"
        f"PARAGRAPHS TO REWRITE:\n{original}"
    )
    result = chat_json(_POLISH_CL_SYSTEM, user, temperature=0.4, max_tokens=700)
    if not isinstance(result, dict):
        return None
    if not all(result.get(k) for k in ("opening_paragraph", "closing_paragraph")):
        return None

    rewritten = "\n\n".join(result.get(k, "") for k in keys)
    if not _fact_guard(original, rewritten):
        return None

    # Strip any salutation the model added back in despite the instruction.
    result["opening_paragraph"] = re.sub(
        r"(?i)^\s*dear[^,\n]*,?\s*", "", result.get("opening_paragraph", "")
    ).lstrip()
    return {k: result.get(k, cl_data.get(k, "")) for k in keys}


def _polish_email(email_data: dict, profile: dict, job: dict,
                  contact: dict) -> Optional[dict]:
    original = f"{email_data.get('subject', '')}\n\n{email_data.get('body', '')}"
    user = (
        f"ROLE: {job.get('title', '')} at {job.get('company', '')}\n"
        f"ADDRESS TO: {contact.get('hr_name') or 'Hiring Manager'}\n\n"
        f"SUBJECT: {email_data.get('subject', '')}\n\n"
        f"BODY:\n{email_data.get('body', '')}"
    )
    result = chat_json(_POLISH_EMAIL_SYSTEM, user, temperature=0.5, max_tokens=700)
    if not isinstance(result, dict):
        return None
    if not (result.get("subject") and result.get("body")):
        return None

    body = result["body"].replace("\\n", "\n")
    if not _fact_guard(original, f"{result['subject']}\n\n{body}"):
        return None

    # The contact block is the candidate's real details — never let a rewrite
    # drop or mangle them.
    if profile.get("email") and profile["email"] not in body:
        logger.warning("[Polish] Rewrite dropped the contact block — keeping original")
        return None

    return {"subject": result["subject"], "body": body}


# ──────────────────────────────────────────────────────────
# CV DOCX
# ──────────────────────────────────────────────────────────

def _write_cv(data: dict, path: str):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(1.8)
        sec.bottom_margin = Cm(1.8)
        sec.left_margin   = Cm(1.8)
        sec.right_margin  = Cm(1.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"  # type: ignore[attr-defined]
    normal.font.size = Pt(10)     # type: ignore[attr-defined]

    _cv_header(doc, data)

    _heading(doc, "PROFILE")
    p = doc.add_paragraph(data.get("profile_summary", ""))
    p.runs[0].font.color.rgb = MID  # type: ignore[attr-defined]
    p.runs[0].font.size = Pt(10)    # type: ignore[attr-defined]
    p.paragraph_format.space_after = Pt(4)

    _heading(doc, "CORE SKILLS")
    for cat, skills in (data.get("core_skills") or {}).items():
        if not skills:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        rc = p.add_run(f"{cat}  ")
        rc.bold = True
        rc.font.size = Pt(9)  # type: ignore[attr-defined]
        rc.font.color.rgb = ACCENT  # type: ignore[attr-defined]
        rs = p.add_run("  ·  ".join(skills))
        rs.font.size = Pt(9)  # type: ignore[attr-defined]
        rs.font.color.rgb = MID  # type: ignore[attr-defined]

    _heading(doc, "PROFESSIONAL EXPERIENCE")
    for exp in (data.get("experience") or []):
        _exp_block(doc, exp)

    _heading(doc, "FEATURED PROJECTS")
    for proj in (data.get("projects") or [])[:4]:
        _proj_block(doc, proj)

    certs = data.get("certifications") or []
    if certs:
        _heading(doc, "CERTIFICATIONS")
        for cert in certs:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Inches(0.2)
            r = p.add_run(cert)
            r.font.color.rgb = MID  # type: ignore[attr-defined]
            r.font.size = Pt(9.5)   # type: ignore[attr-defined]

    _heading(doc, "EDUCATION")
    for edu in (data.get("education") or []):
        p = doc.add_paragraph()
        r1 = p.add_run(edu.get("degree", ""))
        r1.bold = True
        r1.font.color.rgb = DARK  # type: ignore[attr-defined]
        if edu.get("school"):
            r2 = p.add_run(f"  |  {edu['school']}")
            r2.font.color.rgb = MID  # type: ignore[attr-defined]
        p.paragraph_format.space_after = Pt(1)
        if edu.get("year"):
            py = doc.add_paragraph(edu["year"])
            py.runs[0].font.color.rgb = GRAY  # type: ignore[attr-defined]
            py.runs[0].font.size = Pt(9)      # type: ignore[attr-defined]
            py.paragraph_format.space_after = Pt(5)

    doc.save(path)
    logger.info(f"[DocGen] CV saved → {path}")


def _cv_header(doc: Document, data: dict):
    identity = data.get("identity") or {}
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run((identity.get("name") or "").upper())
    r.bold = True
    r.font.size = Pt(22)  # type: ignore[attr-defined]
    r.font.color.rgb = DARK  # type: ignore[attr-defined]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)

    # Titles come from roles actually held, not from a hand-set target list.
    roles = identity.get("titles") or []
    roles_str = "  ·  ".join(roles[:3])
    if roles_str:
        pt = doc.add_paragraph(roles_str)
        pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pt.runs[0].bold = True
        pt.runs[0].font.size = Pt(10)  # type: ignore[attr-defined]
        pt.runs[0].font.color.rgb = ACCENT  # type: ignore[attr-defined]
        pt.paragraph_format.space_after = Pt(4)

    contacts = [c for c in [
        identity.get("phone"), identity.get("email"),
        identity.get("linkedin"), identity.get("github"),
        identity.get("location"),
    ] if c]
    pc = doc.add_paragraph("  ·  ".join(contacts))
    pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pc.runs[0].font.size = Pt(8.5)  # type: ignore[attr-defined]
    pc.runs[0].font.color.rgb = MID  # type: ignore[attr-defined]
    pc.paragraph_format.space_after = Pt(6)
    _bottom_rule(doc, RULE_BLUE, size=12)


def _heading(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    # Left border
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), RULE_BLUE)
    pBdr.append(left)
    pPr.append(pBdr)
    # Left indent
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "120")
    pPr.append(ind)

    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(10.5)  # type: ignore[attr-defined]
    r.font.color.rgb = ACCENT  # type: ignore[attr-defined]
    r.font.all_caps = True     # type: ignore[attr-defined]
    _bottom_rule(doc, RULE_LIGHT, size=4)


def _bottom_rule(doc: Document, color: str, size: int = 4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), str(size))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)


# Phrases that mean "no real data" — filtered out before writing to docx
_PLACEHOLDER_PATTERNS = re.compile(
    r"^\s*(not specified|no details available|n/?a|unknown|none|tbd|placeholder"
    r"|not available|not provided|unspecified)\s*$",
    re.IGNORECASE,
)

def _is_placeholder(value: str) -> bool:
    """Return True if the value is a known placeholder that should be omitted."""
    return not value or bool(_PLACEHOLDER_PATTERNS.match(value.strip()))


def _exp_block(doc: Document, exp: dict):
    title   = exp.get("title", "")
    company = exp.get("company", "")
    period  = exp.get("period", "")
    bullets = [b for b in (exp.get("bullets") or []) if not _is_placeholder(b)]

    # Skip the entire block if there's nothing real to show
    if _is_placeholder(title) and not bullets:
        return

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)

    if not _is_placeholder(title):
        rt = p.add_run(title)
        rt.bold = True
        rt.font.size = Pt(10.5)  # type: ignore[attr-defined]
        rt.font.color.rgb = DARK  # type: ignore[attr-defined]

    if not _is_placeholder(company):
        rs = p.add_run(f"  ·  {company}")
        rs.font.color.rgb = MID  # type: ignore[attr-defined]
        rs.bold = True

    if not _is_placeholder(period):
        rd = p.add_run(f"  |  {period}")
        rd.font.color.rgb = GRAY  # type: ignore[attr-defined]
        rd.font.italic = True  # type: ignore[attr-defined]
        rd.font.size = Pt(9)  # type: ignore[attr-defined]

    for bullet in bullets:
        bp = doc.add_paragraph(style="List Bullet")
        bp.paragraph_format.space_after = Pt(1)
        bp.paragraph_format.left_indent = Inches(0.2)
        rb = bp.add_run(bullet)
        rb.font.color.rgb = MID  # type: ignore[attr-defined]
        rb.font.size = Pt(9.5)  # type: ignore[attr-defined]

    doc.add_paragraph("").paragraph_format.space_after = Pt(2)


def _proj_block(doc: Document, proj: dict):
    name    = proj.get("name", "")
    stack   = proj.get("stack", "")
    bullets = [b for b in (proj.get("bullets") or []) if not _is_placeholder(b)]

    # Skip entirely if nothing real to show
    if _is_placeholder(name):
        return

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    rn = p.add_run(name)
    rn.bold = True
    rn.font.color.rgb = DARK  # type: ignore[attr-defined]
    rn.font.size = Pt(10)  # type: ignore[attr-defined]

    if not _is_placeholder(stack):
        rst = p.add_run(f"  —  {stack}")
        rst.font.color.rgb = ACCENT  # type: ignore[attr-defined]
        rst.font.italic = True  # type: ignore[attr-defined]
        rst.font.size = Pt(9)   # type: ignore[attr-defined]

    # Project URL — shown as plain text (ATS-safe) and clickable in Word
    url = (proj.get("url") or "").strip()
    if url and url.startswith("http"):
        pu = doc.add_paragraph()
        pu.paragraph_format.space_after = Pt(1)
        pu.paragraph_format.left_indent = Inches(0.0)
        ru = pu.add_run(url)
        ru.font.size = Pt(8.5)  # type: ignore[attr-defined]
        ru.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)  # type: ignore[attr-defined]
        ru.font.underline = True  # type: ignore[attr-defined]

    for bullet in bullets:
        bp = doc.add_paragraph(style="List Bullet")
        bp.paragraph_format.space_after = Pt(1)
        bp.paragraph_format.left_indent = Inches(0.2)
        rb = bp.add_run(bullet)
        rb.font.color.rgb = MID  # type: ignore[attr-defined]
        rb.font.size = Pt(9.5)  # type: ignore[attr-defined]

    doc.add_paragraph("").paragraph_format.space_after = Pt(2)


# ──────────────────────────────────────────────────────────
# Cover Letter DOCX
# ──────────────────────────────────────────────────────────

def _write_cl(data: dict, job: dict, contact: dict, path: str,
               identity: dict | None = None):
    identity = identity or {}
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(2.8)
        sec.right_margin  = Cm(2.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"  # type: ignore[attr-defined]
    normal.font.size = Pt(11)     # type: ignore[attr-defined]

    # Header
    ph = doc.add_paragraph()
    rh = ph.add_run((identity.get("name") or "").upper())
    rh.bold = True
    rh.font.size = Pt(16)  # type: ignore[attr-defined]
    rh.font.color.rgb = ACCENT  # type: ignore[attr-defined]
    ph.paragraph_format.space_after = Pt(2)

    contacts = [c for c in [identity.get("email"), identity.get("phone"),
                            identity.get("linkedin")] if c]
    pc = doc.add_paragraph("  |  ".join(contacts))
    pc.runs[0].font.size = Pt(9)  # type: ignore[attr-defined]
    pc.runs[0].font.color.rgb = GRAY  # type: ignore[attr-defined]
    pc.paragraph_format.space_after = Pt(14)
    _bottom_rule(doc, RULE_BLUE)

    doc.add_paragraph(datetime.now().strftime("%B %d, %Y")).paragraph_format.space_after = Pt(14)

    hr_name  = contact.get("hr_name", "")
    hr_title = contact.get("hr_title", "")
    company  = job.get("company", "")
    for line, sp in [(hr_name, 0), (hr_title, 0), (company, 14)]:
        if line:
            px = doc.add_paragraph(line)
            px.paragraph_format.space_after = Pt(sp)

    sal = doc.add_paragraph(f"Dear {hr_name}," if hr_name else "Dear Hiring Manager,")
    sal.paragraph_format.space_after = Pt(12)

    for key in ["opening_paragraph", "body_paragraph_1", "body_paragraph_2", "closing_paragraph"]:
        text = data.get(key, "")
        if not text:
            continue
        # Strip any AI-hallucinated salutation from the start of the opening paragraph
        if key == "opening_paragraph":
            import re as _re
            text = _re.sub(r'(?i)^(dear [^,\n]+,?\s*)', '', text).lstrip()
        if text:
            pp = doc.add_paragraph(text)
            pp.paragraph_format.space_after = Pt(10)

    doc.add_paragraph("Sincerely,").paragraph_format.space_after = Pt(20)
    ps = doc.add_paragraph(identity.get("name", ""))
    ps.runs[0].bold = True
    ps.runs[0].font.color.rgb = DARK  # type: ignore[attr-defined]
    for line in [identity.get("email"), identity.get("phone"),
                 identity.get("linkedin")]:
        if line:
            pl = doc.add_paragraph(line)
            pl.paragraph_format.space_after = Pt(0)
            pl.runs[0].font.color.rgb = GRAY  # type: ignore[attr-defined]

    doc.save(path)
    logger.info(f"[DocGen] Cover letter saved → {path}")


# ──────────────────────────────────────────────────────────
# Email Draft
# ──────────────────────────────────────────────────────────

def _write_email(email_data: dict, job: dict, contact: dict, score_data: dict, folder: Path):
    to_email = (
        contact.get("hr_email")
        or contact.get("application_email")
        or "[ EMAIL NOT FOUND — check application URL ]"
    )
    hr_name  = contact.get("hr_name", "") or "Not found"
    hr_title = contact.get("hr_title", "") or ""
    subject  = email_data.get("subject") or f"Application: {job.get('title','')}"
    body     = email_data.get("body", "").replace("\\n", "\n")
    sep      = "─" * 60

    lines = [
        "╔══════════════════════════════════════════════════════════╗",
        "║                    EMAIL DRAFT                          ║",
        "╚══════════════════════════════════════════════════════════╝",
        "",
        "── SEND TO ───────────────────────────────────────────────",
        f"  TO:       {to_email}",
        f"  HR NAME:  {hr_name}" + (f"  ({hr_title})" if hr_title else ""),
        f"  SUBJECT:  {subject}",
        "",
        "── APPLICATION LINKS ─────────────────────────────────────",
        f"  Job URL:         {job.get('url','—')}",
        f"  Apply URL:       {contact.get('application_url') or job.get('url','—')}",
        f"  App Email (alt): {contact.get('application_email') or '—'}",
        "",
        sep,
        "  EMAIL BODY  (copy from here)",
        sep,
        "",
        body,
        "",
        sep,
        "",
        "── MATCH ANALYSIS ────────────────────────────────────────",
        f"  Score:         {score_data.get('score', 0)}/100",
        f"  Match Reasons: {', '.join(score_data.get('match_reasons', [])) or '—'}",
        f"  Gaps:          {', '.join(score_data.get('gaps', [])) or '—'}",
        f"  ATS Keywords:  {', '.join(score_data.get('ats_keywords', [])) or '—'}",
        f"  Salary Listed: {job.get('salary') or 'Not listed'}",
        f"  Source:        {job.get('source','—')}",
        f"  Generated:     {datetime.now().strftime('%Y-%m-%d %H:%M')}",
    ]
    (folder / "EMAIL_DRAFT.txt").write_text("\n".join(lines), encoding="utf-8")

    # Machine-readable companion for automated senders (the SaaS reads this;
    # the legacy human-driven flow keeps using EMAIL_DRAFT.txt above).
    import json as _json
    (folder / "email.json").write_text(
        _json.dumps({"to": to_email, "subject": subject, "body": body}),
        encoding="utf-8",
    )


# ──────────────────────────────────────────────────────────
# PDF Conversion — Windows aware
# ──────────────────────────────────────────────────────────

def _find_soffice() -> Optional[str]:
    import shutil
    found = shutil.which("soffice")
    if found:
        return found
    if sys.platform == "win32":
        candidates: list[str] = []
        for base in [r"C:\Program Files", r"C:\Program Files (x86)"]:
            bp = Path(base)
            if bp.exists():
                for d in bp.iterdir():
                    if "libreoffice" in d.name.lower():
                        exe = d / "program" / "soffice.exe"
                        if exe.exists():
                            candidates.insert(0, str(exe))
        for c in [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ]:
            if c not in candidates:
                candidates.append(c)
        for c in candidates:
            if Path(c).exists():
                return c
    if sys.platform == "darwin":
        mac = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
        if Path(mac).exists():
            return mac
    return None


def _to_pdf(docx_path: str, out_dir: str):
    pdf_path = Path(out_dir) / (Path(docx_path).stem + ".pdf")

    # 1. docx2pdf
    try:
        from docx2pdf import convert  # type: ignore
        convert(docx_path, str(pdf_path))
        if pdf_path.exists():
            logger.info(f"[PDF] docx2pdf → {pdf_path}")
            return
    except ImportError:
        pass
    except Exception as e:
        logger.warning(f"[PDF] docx2pdf failed: {e}")

    # 2. LibreOffice
    soffice = _find_soffice()
    if soffice:
        try:
            result = subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf", "--outdir", out_dir, docx_path],
                capture_output=True, text=True, timeout=90,
            )
            if result.returncode == 0 and pdf_path.exists():
                logger.info(f"[PDF] LibreOffice → {pdf_path}")
                return
            logger.warning(f"[PDF] LibreOffice stderr: {result.stderr[:200]}")
        except subprocess.TimeoutExpired:
            logger.warning("[PDF] LibreOffice timed out")
        except Exception as e:
            logger.warning(f"[PDF] LibreOffice error: {e}")
    else:
        logger.warning("[PDF] LibreOffice not found in PATH or standard locations. "
                       "Install it OR run: pip install docx2pdf (requires MS Word on Windows)")

    # 3. ReportLab
    _to_pdf_reportlab(docx_path, out_dir)


def _to_pdf_reportlab(docx_path: str, out_dir: str):
    try:
        from reportlab.lib.pagesizes import A4            # type: ignore
        from reportlab.lib.styles import getSampleStyleSheet  # type: ignore
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer  # type: ignore
        from reportlab.lib.units import cm                # type: ignore

        pdf_path = str(Path(out_dir) / (Path(docx_path).stem + ".pdf"))
        texts    = [p.text for p in Document(docx_path).paragraphs if p.text.strip()]
        styles   = getSampleStyleSheet()
        story    = []
        for t in texts:
            story.append(Paragraph(t.replace("&", "&amp;").replace("<", "&lt;"), styles["Normal"]))
            story.append(Spacer(1, 0.15 * cm))
        pdf = SimpleDocTemplate(pdf_path, pagesize=A4,
                                leftMargin=2*cm, rightMargin=2*cm,
                                topMargin=2*cm, bottomMargin=2*cm)
        pdf.build(story)
        logger.info(f"[PDF] ReportLab → {pdf_path}")
    except Exception as e:
        logger.error(f"[PDF] All methods failed: {e}")


# ──────────────────────────────────────────────────────────
# Utility
# ──────────────────────────────────────────────────────────

def _safe_dirname(name: str) -> str:
    name = re.sub(r"[^\w\s-]", "", name)
    name = re.sub(r"\s+", "_", name.strip())
    return name[:50]
